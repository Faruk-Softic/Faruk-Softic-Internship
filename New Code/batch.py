"""
Grades papers holistically using the OpenAI Batch API.
Submits all runs as a single batch job, polls until complete,
then parses results into the same CSV format as the synchronous version.
Checklist items (0/1), section grades, and feedback are all written to CSV.
Final grade is computed in Python only — never by the LLM.
"""

import os, csv, json, glob, time, io
from datetime import datetime
from typing import Optional

import pdfplumber, tiktoken
from docx import Document
from openai import OpenAI
from dotenv import load_dotenv

load_dotenv()

# ── Paths & constants ─────────────────────────────────────────────────────────

PAPERS_FOLDER        = os.environ.get("PAPERS_FOLDER")
RUBRIC_ORIGINAL_PATH = os.environ.get("RUBRIC_ORIGINAL_PATH")
RUBRIC_IMPROVED_PATH = os.environ.get("RUBRIC_IMPROVED_PATH")
WRITING_GUIDE        = os.environ.get("WRITING_GUIDE")
SAMPLE_RESULTS       = os.environ.get("SAMPLE_RESULTS")
CALIBRATION_SUMMARY  = os.environ.get("CALIBRATION_SUMMARY")
PIPELINES_FILE       = os.environ.get("PIPELINES_FILE")
RESULTS_CSV          = os.environ.get("RESULTS_CSV")
PROMPTS_FILE         = os.environ.get("PROMPTS_FILE")

N_RUNS        = 3
POLL_INTERVAL = 60  # seconds between status checks

WEIGHTS = {
    "introduction":   0.30,
    "methods":        0.15,
    "results":        0.15,
    "discussion":     0.30,
    "language_style": 0.10,
}

TIER_ORDER  = ["tier0", "tier1", "tier2", "tier3", "tier4"]
ITEM_TYPES  = {
    "tier0": ["ko"],
    "tier1": ["core_criteria", "additional_criteria"],
    "tier2": ["core_criteria", "additional_criteria"],
    "tier3": ["core_criteria", "additional_criteria"],
    "tier4": ["core_criteria", "additional_criteria"],
}
TYPE_ABBREV = {"ko": "ko", "core_criteria": "cc", "additional_criteria": "ac"}
TYPE_LABELS = {
    "ko":                  "KNOCK-OUT — paper fails to reach Tier 1 if ANY apply:",
    "core_criteria":       "CORE CRITERIA — all must be met to reach this tier:",
    "additional_criteria": "ADDITIONAL CRITERIA — push grade toward top of tier:",
}
TIER_LABELS = {
    "tier0": "Tier 0 — Insufficient (1.0–5.0)",
    "tier1": "Tier 1 — Passing (5.5–6.4)",
    "tier2": "Tier 2 — Adequate (6.5–7.4)",
    "tier3": "Tier 3 — Good (7.5–8.4)",
    "tier4": "Tier 4 — Excellent (8.5–10.0)",
}
SECTION_LABELS = {
    "introduction":   "Introduction",
    "methods":        "Methods",
    "results":        "Results",
    "discussion":     "Discussion",
    "language_style": "Language & Style",
}

client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))


# ── Prompt loading & building ─────────────────────────────────────────────────

def load_prompts(path: str) -> dict:
    """Load all prompt strings and checklist data from prompts.json."""
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def get_checklist_keys(checklist: dict) -> list[str]:
    """
    Return an ordered list of all checklist column names, e.g.:
      introduction_t0_ko_1, introduction_t1_cc_1, introduction_t1_ac_1, ...
      language_style_t4_ac_2
    """
    keys = []
    for section, tiers in checklist.items():
        for tier in TIER_ORDER:
            if tier not in tiers:
                continue
            t = tier.replace("tier", "t")
            for item_type in ITEM_TYPES[tier]:
                if item_type not in tiers[tier]:
                    continue
                abbrev = TYPE_ABBREV[item_type]
                for idx in range(1, len(tiers[tier][item_type]) + 1):
                    keys.append(f"{section}_{t}_{abbrev}_{idx}")
    return keys


def build_rubric_text(checklist: dict) -> str:
    """Render the structured checklist into the plain-text rubric the LLM sees."""
    lines = [
        "Tiers are cumulative. A paper must satisfy all CORE CRITERIA items in a tier "
        "(and all tiers below it) to be graded within that tier's range. "
        "ADDITIONAL CRITERIA items push the grade toward the top of the range but are not required.\n"
        "Tier 0 items are knock-out criteria: if ANY apply, the paper cannot reach Tier 1.\n\n"
        "Tier ranges:\n"
        "  Tier 0 — Insufficient  : 1.0 – 5.0\n"
        "  Tier 1 — Passing       : 5.5 – 6.4\n"
        "  Tier 2 — Adequate      : 6.5 – 7.4\n"
        "  Tier 3 — Good          : 7.5 – 8.4\n"
        "  Tier 4 — Excellent     : 8.5 – 10.0"
    ]
    for section, tiers in checklist.items():
        sec_label = SECTION_LABELS.get(section, section)
        lines += [f"\n{'=' * 60}", f"SECTION: {sec_label.upper()}", "=" * 60]
        for tier in TIER_ORDER:
            if tier not in tiers:
                continue
            lines += [f"\n{TIER_LABELS[tier]}", "-" * 40]
            for item_type in ITEM_TYPES[tier]:
                if item_type not in tiers[tier]:
                    continue
                lines.append(TYPE_LABELS[item_type])
                t      = tier.replace("tier", "t")
                abbrev = TYPE_ABBREV[item_type]
                for idx, text in enumerate(tiers[tier][item_type], start=1):
                    key = f"{section}_{t}_{abbrev}_{idx}"
                    lines.append(f"  [{key}] {text}")
    return "\n".join(lines)


def build_output_format(checklist: dict, prompts: dict) -> str:
    """
    Generate the full output-format specification shown to the LLM,
    with every checklist key and all grade/feedback keys listed explicitly.
    """
    lines = [prompts["output_format_preamble"], "", "{"]
    for section, tiers in checklist.items():
        lines.append(f'  // --- {SECTION_LABELS.get(section, section)} checklist ---')
        for tier in TIER_ORDER:
            if tier not in tiers:
                continue
            t = tier.replace("tier", "t")
            for item_type in ITEM_TYPES[tier]:
                if item_type not in tiers[tier]:
                    continue
                abbrev = TYPE_ABBREV[item_type]
                for idx in range(1, len(tiers[tier][item_type]) + 1):
                    key = f"{section}_{t}_{abbrev}_{idx}"
                    lines.append(f'  "{key}": <0 or 1>,')
        lines.append("")

    grade_feedback_keys = [
        ("introduction",   "intro_grade",      "intro_feedback"),
        ("methods",        "methods_grade",     "methods_feedback"),
        ("results",        "results_grade",     "results_feedback"),
        ("discussion",     "discussion_grade",  "discussion_feedback"),
        ("language_style", "lang_style_grade",  "lang_style_feedback"),
    ]
    lines.append("  // --- Section grades and feedback ---")
    for section, grade_key, feedback_key in grade_feedback_keys:
        label = SECTION_LABELS.get(section, section)
        lines.append(f'  "{grade_key}": <number 1.0-10.0 in 0.5 steps, or null if section absent>,')
        lines.append(f'  "{feedback_key}": "<reasoning behind the {label} grade>",')

    lines[-1] = lines[-1].rstrip(",")
    lines.append("}")
    return "\n".join(lines)


def build_system_prompt(rubric_text: str, resources: dict, prompts: dict,
                        output_format: str) -> str:
    return prompts["system_prompt_template"].format(
        llm_role            = prompts["llm_role"],
        resources_preamble  = prompts["resources_preamble"],
        writing_guide       = resources["writing_guide"],
        sample_results      = resources["sample_results"],
        calibration_summary = resources["calibration_summary"],
        grading_notes       = prompts["grading_notes"],
        rubric              = rubric_text,
        output_format       = output_format,
    ).strip()


def build_user_message(paper_text: str, prompts: dict) -> str:
    return prompts["user_message_prefix"] + paper_text + prompts["user_message_suffix"]


# ── File readers ──────────────────────────────────────────────────────────────

def _read_pdf(path: str) -> str:
    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
            for table in page.extract_tables() or []:
                for row in table:
                    row_text = " | ".join(cell.strip() for cell in row if cell and cell.strip())
                    if row_text:
                        parts.append(row_text)
    return "\n\n".join(parts)


def _read_docx(path: str) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _read_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def load_document(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":  return _read_pdf(path)
    if ext == ".docx": return _read_docx(path)
    if ext == ".txt":  return _read_txt(path)
    raise ValueError(f"Unsupported format '{ext}'. Use .pdf, .docx, or .txt.")


# ── Token counting ────────────────────────────────────────────────────────────

_ENC        = tiktoken.get_encoding("o200k_base")
_MAX_TOKENS = 100_000


def check_token_count(text: str, label: str = "") -> int:
    n   = len(_ENC.encode(text))
    tag = f"[{label}] " if label else ""
    if n > _MAX_TOKENS:
        raise ValueError(f"{tag}{n} tokens exceeds the {_MAX_TOKENS}-token limit.")
    if n > _MAX_TOKENS * 0.85:
        print(f"  ⚠  Warning: {tag}{n} tokens (>85% of limit).")
    return n


# ── Resource & rubric loading ─────────────────────────────────────────────────

def load_rubrics() -> dict[str, str]:
    print("\nLoading rubrics...")
    rubrics = {}
    for name, path in [("original", RUBRIC_ORIGINAL_PATH), ("improved", RUBRIC_IMPROVED_PATH)]:
        if not os.path.exists(path):
            raise FileNotFoundError(f"Rubric not found: {path}")
        rubrics[name] = _read_docx(path)
        print(f"  ✓ '{name}' ({check_token_count(rubrics[name], label=name)} tokens)")
    return rubrics


def load_resources() -> dict[str, str]:
    print("Loading resources...")
    resources = {}
    for label, path, reader in [
        ("writing_guide",       WRITING_GUIDE,       _read_pdf),
        ("sample_results",      SAMPLE_RESULTS,      _read_pdf),
        ("calibration_summary", CALIBRATION_SUMMARY, _read_docx),
    ]:
        if not os.path.exists(path):
            raise FileNotFoundError(f"Resource not found: {path}")
        resources[label] = reader(path)
        print(f"  ✓ '{label}' ({check_token_count(resources[label], label=label)} tokens)")
    return resources


# ── Pipeline & paper discovery ────────────────────────────────────────────────

def load_pipelines(path: str) -> list[dict]:
    with open(path, "r", encoding="utf-8") as f:
        all_pipelines = json.load(f)
    holistic = [p for p in all_pipelines if p.get("grading_mode") == "holistic"]
    print(f"Loaded {len(holistic)} holistic pipeline(s) from {len(all_pipelines)} total.")
    return holistic


def list_papers(folder: str) -> list[tuple[str, str]]:
    paths = (
        glob.glob(os.path.join(folder, "*.pdf")) +
        glob.glob(os.path.join(folder, "*.docx")) +
        glob.glob(os.path.join(folder, "*.txt"))
    )
    papers = [(os.path.splitext(os.path.basename(p))[0], p)
              for p in paths if not os.path.basename(p).startswith("~$")]
    return sorted(papers)


# ── Grade parsing & final grade computation ───────────────────────────────────

def _parse_grade(value) -> Optional[float]:
    if value is None:
        return None
    try:
        g = float(value)
    except (TypeError, ValueError):
        return None
    if g < 1.0 or g > 10.0:
        return None
    return round(g * 2) / 2.0


def _parse_binary(value) -> Optional[int]:
    """Accept 0 or 1 only. Returns None if value is missing or invalid."""
    if value is None:
        return None
    try:
        v = int(value)
    except (TypeError, ValueError):
        return None
    return v if v in (0, 1) else None


def compute_final_grade(grades: dict[str, Optional[float]]) -> Optional[float]:
    if any(grades.get(k) is None for k in WEIGHTS):
        return None
    return sum(grades[k] * w for k, w in WEIGHTS.items())


def parse_reply(raw: str, checklist_keys: list[str]) -> dict:
    """Parse the LLM JSON reply into checklist 0/1 values, grades, and feedback."""
    result = {k: None for k in checklist_keys}
    for k in ["intro_grade", "methods_grade", "results_grade", "discussion_grade",
              "lang_style_grade", "intro_feedback", "methods_feedback",
              "results_feedback", "discussion_feedback", "lang_style_feedback"]:
        result[k] = None
    result["raw_reply"]        = raw
    result["missing_sections"] = ""

    try:
        data = json.loads(raw)
    except Exception:
        return result

    for key in checklist_keys:
        result[key] = _parse_binary(data.get(key))

    for k in ["intro_grade", "methods_grade", "results_grade",
              "discussion_grade", "lang_style_grade"]:
        result[k] = _parse_grade(data.get(k))

    for k in ["intro_feedback", "methods_feedback", "results_feedback",
              "discussion_feedback", "lang_style_feedback"]:
        result[k] = data.get(k)

    missing = [
        sec for sec, gk in {
            "introduction":   "intro_grade",
            "methods":        "methods_grade",
            "results":        "results_grade",
            "discussion":     "discussion_grade",
            "language_style": "lang_style_grade",
        }.items() if result[gk] is None
    ]
    result["missing_sections"] = ", ".join(missing) if missing else ""
    return result


# ── Batch request building ────────────────────────────────────────────────────

def build_batch_request(custom_id: str, system_prompt: str, paper_text: str,
                        model: str, temperature: float, prompts: dict) -> dict:
    """Build a single JSONL request object for the Batch API."""
    return {
        "custom_id": custom_id,
        "method":    "POST",
        "url":       "/v1/responses",
        "body": {
            "model":        model,
            "temperature":  temperature,
            "instructions": system_prompt,
            "input":        build_user_message(paper_text, prompts),
            "text":         {"format": {"type": "json_object"}},
        },
    }


# ── CSV writing ───────────────────────────────────────────────────────────────

def append_result_row(csv_path: str, row: dict) -> None:
    file_exists = os.path.exists(csv_path)
    os.makedirs(os.path.dirname(csv_path), exist_ok=True)
    with open(csv_path, "a", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=row.keys())
        if not file_exists:
            writer.writeheader()
        writer.writerow(row)


# ── Result parsing ────────────────────────────────────────────────────────────

def process_batch_result(result_line: dict, metadata: dict,
                         checklist_keys: list[str]) -> dict | None:
    """
    Parse one line from the batch output file into a CSV-ready row.
    Returns None and prints an error message if the request failed.
    """
    custom_id = result_line.get("custom_id", "unknown")
    meta      = metadata.get(custom_id)

    if meta is None:
        print(f"  ⚠  No metadata found for custom_id '{custom_id}' — skipping.")
        return None

    error = result_line.get("error")
    if error:
        print(f"  ✗ API error for '{custom_id}': {error}")
        return None

    try:
        raw_reply = result_line["response"]["body"]["output"][0]["content"][0]["text"]
    except (KeyError, IndexError, TypeError) as e:
        print(f"  ✗ Could not extract reply for '{custom_id}': {e}")
        print(f"      Raw line: {json.dumps(result_line)[:300]}")
        return None

    parsed = parse_reply(raw_reply, checklist_keys)
    grades = {
        "introduction":   parsed["intro_grade"],
        "methods":        parsed["methods_grade"],
        "results":        parsed["results_grade"],
        "discussion":     parsed["discussion_grade"],
        "language_style": parsed["lang_style_grade"],
    }
    final_grade = compute_final_grade(grades)

    if parsed["missing_sections"]:
        print(f"  ⚠  '{custom_id}' — missing section(s): {parsed['missing_sections']}")
    if final_grade is None:
        print(f"  ⚠  '{custom_id}' — could not compute final grade.")

    # ── CSV row: metadata | checklist 0/1 | grades + feedback | final ─────────
    row = {
        "run_id":         meta["run_id"],
        "student_id":     meta["student_id"],
        "grading_mode":   "holistic",
        "rubric_version": meta["rubric_version"],
        "model":          meta["model"],
        "temperature":    meta["temperature"],
        "timestamp":      datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }
    for key in checklist_keys:
        row[key] = parsed[key]
    row.update({
        "intro_grade":         parsed["intro_grade"],
        "intro_feedback":      parsed["intro_feedback"],
        "methods_grade":       parsed["methods_grade"],
        "methods_feedback":    parsed["methods_feedback"],
        "results_grade":       parsed["results_grade"],
        "results_feedback":    parsed["results_feedback"],
        "discussion_grade":    parsed["discussion_grade"],
        "discussion_feedback": parsed["discussion_feedback"],
        "lang_style_grade":    parsed["lang_style_grade"],
        "lang_style_feedback": parsed["lang_style_feedback"],
        "final_grade":         final_grade,
        "missing_sections":    parsed["missing_sections"],
    })
    return row


# ── Entry point ───────────────────────────────────────────────────────────────

if __name__ == "__main__":

    start_time = time.time()

    PROMPTS   = load_prompts(PROMPTS_FILE)
    RUBRICS   = load_rubrics()
    RESOURCES = load_resources()
    PIPELINES = load_pipelines(PIPELINES_FILE)

    checklist      = PROMPTS["checklist"]
    CHECKLIST_KEYS = get_checklist_keys(checklist)
    RUBRIC_TEXT    = build_rubric_text(checklist)
    OUTPUT_FORMAT  = build_output_format(checklist, PROMPTS)

    print(f"Checklist columns : {len(CHECKLIST_KEYS)}")

    # Resume support — skip already completed run_ids
    completed: set[str] = set()
    if os.path.exists(RESULTS_CSV):
        with open(RESULTS_CSV, "r", encoding="utf-8") as f:
            completed = {row["run_id"] for row in csv.DictReader(f) if "run_id" in row}
    print(f"Skipping {len(completed)} already completed run(s).")

    papers = list_papers(PAPERS_FOLDER)
    print(f"Found {len(papers)} paper(s): {[sid for sid, _ in papers]}\n")

    # ── Build JSONL batch + metadata ──────────────────────────────────────────

    batch_requests = []
    metadata       = {}
    skipped_papers = []

    print("Preparing batch requests...")
    for student_id, doc_path in papers:
        try:
            paper_text = load_document(doc_path)
            check_token_count(paper_text, label=student_id)
            print(f"  ✓ '{student_id}': {len(paper_text)} characters loaded.")
        except Exception as e:
            print(f"  ✗ Cannot load '{student_id}': {e} — skipping.")
            skipped_papers.append(student_id)
            continue

        for pipeline in PIPELINES:
            rubric_version = pipeline["rubric_version"]
            model          = pipeline["model"]
            temperature    = pipeline.get("temperature", 1.0)

            # Improved rubric: auto-generated from structured checklist.
            # Original rubric: loaded from file as before.
            rubric_text   = (RUBRIC_TEXT if rubric_version == "improved"
                             else RUBRICS[rubric_version])
            system_prompt = build_system_prompt(rubric_text, RESOURCES, PROMPTS, OUTPUT_FORMAT)

            for run_idx in range(1, N_RUNS + 1):
                run_id    = f"{student_id}.{pipeline['pipeline_id']}.run{run_idx}"
                custom_id = run_id

                if run_id in completed:
                    print(f"  -> Skipping {run_id} (already done).")
                    continue

                batch_requests.append(
                    build_batch_request(custom_id, system_prompt, paper_text,
                                        model, temperature, PROMPTS)
                )
                metadata[custom_id] = {
                    "run_id":         run_id,
                    "student_id":     student_id,
                    "rubric_version": rubric_version,
                    "model":          model,
                    "temperature":    temperature,
                }

    total_requests = len(batch_requests)
    print(f"\n{'=' * 64}")
    print(f"  Papers found       : {len(papers)}")
    print(f"  Papers skipped     : {len(skipped_papers)}" + (f" {skipped_papers}" if skipped_papers else ""))
    print(f"  Runs already done  : {len(completed)}")
    print(f"  Requests to submit : {total_requests}")
    print(f"{'=' * 64}\n")

    if total_requests == 0:
        print("Nothing to submit. Exiting.")
        exit(0)

    # ── Upload JSONL file ─────────────────────────────────────────────────────

    print("Uploading batch input file...")
    jsonl_bytes = "\n".join(json.dumps(r) for r in batch_requests).encode("utf-8")
    batch_file  = client.files.create(
        file=("batch_input.jsonl", io.BytesIO(jsonl_bytes), "application/jsonl"),
        purpose="batch",
    )
    print(f"  ✓ Uploaded — file_id: {batch_file.id}")

    # ── Submit batch job ──────────────────────────────────────────────────────

    print("Submitting batch job...")
    batch_job = client.batches.create(
        input_file_id=batch_file.id,
        endpoint="/v1/responses",
        completion_window="24h",
    )
    print(f"  ✓ Submitted — batch_id: {batch_job.id}")
    print(f"  Status: {batch_job.status}\n")

    # ── Poll until complete ───────────────────────────────────────────────────

    print("Polling for completion (press Ctrl+C to abort)...")
    print(f"{'─' * 64}")

    while True:
        time.sleep(POLL_INTERVAL)
        batch_job = client.batches.retrieve(batch_job.id)
        elapsed   = time.strftime("%H:%M:%S", time.gmtime(time.time() - start_time))
        counts    = batch_job.request_counts

        print(
            f"  [{elapsed}]  status={batch_job.status}  |  "
            f"total={counts.total}  completed={counts.completed}  failed={counts.failed}"
        )

        if batch_job.status in ("completed", "failed", "expired", "cancelled"):
            break

    print(f"{'─' * 64}")
    elapsed_final = time.strftime("%H:%M:%S", time.gmtime(time.time() - start_time))
    print(f"  Batch finished — status: {batch_job.status}  |  total elapsed: {elapsed_final}\n")

    if batch_job.status != "completed":
        print(f"  ✗ Batch did not complete successfully (status: {batch_job.status}).")
        print(f"    Check the OpenAI dashboard for batch_id: {batch_job.id}")
        exit(1)

    # ── Download & parse results ──────────────────────────────────────────────

    print("Downloading results...")
    result_content = client.files.content(batch_job.output_file_id).text
    result_lines   = [json.loads(line) for line in result_content.strip().splitlines()]
    print(f"  ✓ Retrieved {len(result_lines)} result line(s).\n")

    print("Parsing and saving results...")
    saved  = 0
    errors = 0
    for line in result_lines:
        row = process_batch_result(line, metadata, CHECKLIST_KEYS)
        if row is not None:
            append_result_row(RESULTS_CSV, row)
            saved += 1
        else:
            errors += 1

    # ── Summary ───────────────────────────────────────────────────────────────

    print(f"\n{'=' * 64}")
    print(f"  Batch grading complete.")
    print(f"  Saved   : {saved} row(s)")
    print(f"  Errors  : {errors} row(s)")
    if errors > 0:
        print(f"  ⚠  Check error messages above for details on failed requests.")
    print(f"  Output  : {RESULTS_CSV}")
    print(f"  Elapsed : {elapsed_final}")
    print(f"{'=' * 64}")
