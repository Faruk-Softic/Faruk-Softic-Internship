"""
This script builds the prompt(s) for batch_run.py. It reads the rubric, writing guide, sample results, 
and calibration summary from the specified files, loads the student papers from the specified folder, 
and generates a JSONL file containing one request per paper per pipeline per run. 

Outputs:
  - prompts_example.json   : one request per pipeline for a single paper (for inspection)
  - batch_input.jsonl      : all requests for all papers x pipelines x runs (for batch submission)
"""

import os
import json
import glob
from dotenv import load_dotenv
import pdfplumber
from docx import Document

load_dotenv()

# Paths

PAPERS_FOLDER        = os.environ.get("PAPERS_FOLDER")
RUBRIC_ORIGINAL_PATH = os.environ.get("RUBRIC_ORIGINAL_PATH")
RUBRIC_IMPROVED_PATH = os.environ.get("RUBRIC_IMPROVED_PATH")
WRITING_GUIDE        = os.environ.get("WRITING_GUIDE")
SAMPLE_RESULTS       = os.environ.get("SAMPLE_RESULTS")
CALIBRATION_SUMMARY  = os.environ.get("CALIBRATION_SUMMARY")
PIPELINES_FILE       = os.environ.get("PIPELINES_FILE")
EXAMPLE_JSON         = os.environ.get("EXAMPLE_JSON", "prompts_example.json")
BATCH_JSONL          = os.environ.get("BATCH_JSONL", "batch_input.jsonl")
EXAMPLE_PAPER_ID     = os.environ.get("EXAMPLE_PAPER_ID")  # optional: pin a specific paper

N_RUNS = 3

LLM_ROLE = (
    "You are an independent grader tasked with grading first scientific papers written by second-year psychology "
    "students. The course focuses on scientific reasoning and argumentation. Keep in mind that "
    "these students have not yet learned how to write a full paper (that is the purpose of "
    "the assignment that you are about to grade) and have only conducted a "
    "simple literature review. Grade against the provided rubric, "
    "writing guide, calibration summary, and sample paper — not against publishable-paper standards. Therefore, "
    "it is possible for a student paper to receive a high grade, even if it does not meet publishable-paper standards. "
    "An abstract is not required, and can therefore never affect the grade negatively, even if it is poorly written or absent. "
    "However, if it is very well-written, it can positively affect the grade. For example, if you "
    "are on the fence between a 7.5 and 8, a strong abstract could tip the scale towards 8. "
    "As a general rule, a paper graded between 6.5 and 7.5 is around average. Grades 8-9 are considered "
    "above average; grades 9 and up are considered excellent. Grades lower than 6.5 are considered "
    "below average. The passing grade is 5.5. The paper you are about to grade has already been graded by "
    "a previous grader, and it belongs to a database of papers that have received grades ranging from 5.0 to 9.0. Your "
    "task is to provide as objective and accurate a grade as possible given the resources and rubrics of the course, in "
    "order to examine how consistent the grades for each paper are across different graders. "
    "Keep in mind that it is essential that you are as objective as "
    "possible. If you think a student deserves a failing grade in the introduction, and you give them "
    "a passing grade, that subgrade might be enough to allow them to pass the course, and you might thus "
    "be setting them up for failure in the next year where they will be expected to work on a more serious "
    "paper with greater autonomy. If you think a student deserves a passing grade, and you give them a failing "
    "grade, you are setting them back a year for no reason. In either case, it is important to be as objective "
    "as possible and to use the rubric and other resources to guide your grading decisions. "
)

# Checklist items imported from config.py.
# These mirror the rubric exactly. The LLM returns true/false for each.

from config import SECTION_KEYS, CHECKLIST

# File-loading functions

def load_document(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        parts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                if text := page.extract_text():
                    parts.append(text)
                for table in page.extract_tables() or []:
                    for row in table:
                        if row_text := " | ".join(c.strip() for c in row if c and c.strip()):
                            parts.append(row_text)
        return "\n\n".join(parts)
    if ext == ".docx":
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                if row_text := " | ".join(c.text.strip() for c in row.cells if c.text.strip()):
                    parts.append(row_text)
        return "\n".join(parts)
    if ext == ".txt":
        return open(path, encoding="utf-8").read()
    raise ValueError(f"Unsupported format '{ext}'.")


def load_files(specs):
    result = {}
    for label, path in specs:
        if not os.path.exists(path):
            raise FileNotFoundError(f"File not found: {path}")
        result[label] = load_document(path)
        print(f"  ✓ '{label}' loaded.")
    return result


def load_pipelines(path):
    include_feedback = os.environ.get("INCLUDE_FEEDBACK", "true").lower() == "true"
    pipelines = [p for p in json.load(open(path, encoding="utf-8"))
                 if p.get("grading_mode") == "holistic"
                 and p.get("include_feedback") == include_feedback]
    print(f"  ✓ {len(pipelines)} holistic pipeline(s) loaded (include_feedback={include_feedback}).")
    return pipelines


def list_papers(folder):
    paths = (glob.glob(os.path.join(folder, "*.pdf")) +
             glob.glob(os.path.join(folder, "*.docx")) +
             glob.glob(os.path.join(folder, "*.txt")))
    return sorted(
        (os.path.splitext(os.path.basename(p))[0], p)
        for p in paths if not os.path.basename(p).startswith("~$")
    )

# Output format builder

def build_output_format(include_feedback: bool) -> str:
    sections = ["intro", "methods", "results", "discussion", "lang_style"]

    checklist_lines = []
    for section in sections:
        for key in CHECKLIST[section]:
            checklist_lines.append(f'  "{key}": <true if criterion is met, false if not>')

    grade_lines = "\n".join(
        f'  "{s}_grade": <number 1.0-10.0 in 0.5 steps, or null if section is absent>,'
        for s in sections
    )

    checklist_block = "\n".join(checklist_lines)

    feedback_block = ""
    if include_feedback:
        feedback_block = "\n" + "\n".join(
            f'  "{s}_feedback": "<reasoning behind the {s.replace("_", " ").title()} grade>",'
            for s in sections
        )

    return f"""Respond with a JSON object containing EXACTLY these keys:
{{
{grade_lines}
{checklist_block}{feedback_block}
}}

If a section appears to be entirely absent from the paper, set its grade to null \
and set all checklist items for that section to false."""


# Prompt builder

def build_system_prompt(rubric_text: str, resources: dict, include_feedback: bool) -> str:
    return f"""{LLM_ROLE}

## Resources
The rubric is your primary grading tool — all grades must be grounded in its criteria. \
The other resources provide supporting context and do not override the rubric.

### Writing guide
{resources['writing_guide']}

### Sample Results section
{resources['sample_results']}

### Calibration summary
{resources['calibration_summary']}

## Grading rubric
{rubric_text}

## Output format
{build_output_format(include_feedback)}""".strip()


def build_request(custom_id: str, system_prompt: str, paper_text: str,
                  model: str, temperature: float) -> dict:
    return {
        "custom_id": custom_id,
        "method": "POST",
        "url": "/v1/responses",
        "body": {
            "model": model,
            "temperature": temperature,
            "instructions": system_prompt,
            "input": "Please grade the following student paper.\n\n" + paper_text + "\n\nRespond with valid JSON.",
            "text": {"format": {"type": "json_object"}},
        },
    }


# Main execution

if __name__ == "__main__":

    print("\nLoading rubrics...")
    RUBRICS = load_files([("original", RUBRIC_ORIGINAL_PATH), ("improved", RUBRIC_IMPROVED_PATH)])

    print("Loading resources...")
    RESOURCES = load_files([
        ("writing_guide",       WRITING_GUIDE),
        ("sample_results",      SAMPLE_RESULTS),
        ("calibration_summary", CALIBRATION_SUMMARY),
    ])

    print("Loading pipelines...")
    PIPELINES = load_pipelines(PIPELINES_FILE)
    print(f"  Reading pipelines from: {PIPELINES_FILE}")

    papers = list_papers(PAPERS_FOLDER)
    print(f"  ✓ {len(papers)} paper(s) found.\n")

    # Build final JSONL for batch submission

    all_requests = []
    for student_id, doc_path in papers:
        try:
            paper_text = load_document(doc_path)
        except Exception as e:
            print(f"  ✗ Cannot load '{student_id}': {e} — skipping.")
            continue

        for pipeline in PIPELINES:
            system_prompt = build_system_prompt(
                RUBRICS[pipeline["rubric_version"]],
                RESOURCES,
                pipeline["include_feedback"],
            )
            for run_idx in range(1, N_RUNS + 1):
                custom_id = f"{student_id}.{pipeline['pipeline_id']}.run{run_idx}"
                all_requests.append(
                    build_request(custom_id, system_prompt, paper_text,
                                  pipeline["model"], pipeline.get("temperature", 1.0))
                )

    with open(BATCH_JSONL, "w", encoding="utf-8") as f:
        for req in all_requests:
            f.write(json.dumps(req) + "\n")
    print(f"  ✓ JSONL written: {BATCH_JSONL} ({len(all_requests)} requests)")

    # Create an example JSON

    if EXAMPLE_PAPER_ID:
        example_paper = next(((sid, p) for sid, p in papers if sid == EXAMPLE_PAPER_ID), None)
        if example_paper is None:
            print(f"  ⚠  EXAMPLE_PAPER_ID '{EXAMPLE_PAPER_ID}' not found — using first paper.")
            example_paper = papers[0]
    else:
        example_paper = papers[0]

    example_student_id, example_path = example_paper
    example_text = load_document(example_path)
    example_entries = []

    for pipeline in PIPELINES:
        system_prompt = build_system_prompt(
            RUBRICS[pipeline["rubric_version"]],
            RESOURCES,
            pipeline["include_feedback"],
        )
        custom_id = f"{example_student_id}.{pipeline['pipeline_id']}.run1"
        example_entries.append(
            build_request(custom_id, system_prompt, example_text,
                          pipeline["model"], pipeline.get("temperature", 1.0))
        )

    with open(EXAMPLE_JSON, "w", encoding="utf-8") as f:
        json.dump(example_entries, f, indent=2, ensure_ascii=False)
    print(f"  ✓ Example JSON written: {EXAMPLE_JSON} ({len(example_entries)} entries, paper: '{example_student_id}')")