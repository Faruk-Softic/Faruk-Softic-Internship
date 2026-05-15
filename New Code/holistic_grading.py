"""
Grades papers holistically in a single API call per run.
Both rubric conditions (original / improved) produce identical output structure.
Final grade is computed in Python only — never by the LLM.
"""

import os
import csv
import json
import glob
import re
from datetime import datetime
from typing import Optional

import pdfplumber
import tiktoken
from docx import Document
from openai import OpenAI


# ── Paths & constants ─────────────────────────────────────────────────────────

API_KEY_FILE         = r"C:\Users\fsoft\Desktop\api_key.txt"
PAPERS_FOLDER        = r"C:\Users\fsoft\Desktop\Za Internship - code\Papers"
RUBRIC_ORIGINAL_PATH = r"C:\Users\fsoft\Desktop\Za Internship - code\Rubric_original.docx"
RUBRIC_IMPROVED_PATH = r"C:\Users\fsoft\Desktop\Za Internship - code\Rubric_improved.docx"
WRITING_GUIDE        = r"C:\Users\fsoft\Desktop\Za Internship - code\Writing_guide.pdf"
SAMPLE_RESULTS       = r"C:\Users\fsoft\Desktop\Za Internship - code\Sample_results.pdf"
CALIBRATION_SUMMARY  = r"C:\Users\fsoft\Desktop\Za Internship - code\Calibration_summary.docx"
PIPELINES_FILE       = r"C:\Users\fsoft\Desktop\Za Internship - code\pipelines.json"
RESULTS_CSV          = r"C:\Users\fsoft\Desktop\Za Internship - code\results.csv"

N_RUNS        = 5
SECTION_ORDER = ["introduction", "methods", "results", "discussion"]

WEIGHTS = {
    "introduction":   0.30,
    "methods":        0.15,
    "results":        0.15,
    "discussion":     0.30,
    "language_style": 0.10,
}

LLM_ROLE = (
    "You are a tutor grading first scientific papers written by second-year psychology "
    "students. The course focuses on scientific reasoning and argumentation. Students have "
    "only previously written a short literature review. Grade against the provided rubric, "
    "writing guide, and sample paper — not against publishable-paper standards. When "
    "deciding whether to fail a student, ask yourself whether they would be able to pass "
    "a bachelor's thesis the following year with the same skills."
)


# ── API client ────────────────────────────────────────────────────────────────

def _load_api_key() -> str:
    if not os.path.exists(API_KEY_FILE):
        raise FileNotFoundError(f"API key file not found: {API_KEY_FILE}")
    with open(API_KEY_FILE, "r", encoding="utf-8") as f:
        key = f.readline().strip()
    if not key:
        raise ValueError(f"API key file is empty: {API_KEY_FILE}")
    return key

client = OpenAI(api_key=_load_api_key())


# ── File readers ──────────────────────────────────────────────────────────────

def _read_pdf(path: str) -> str:
    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
            for table in page.extract_tables() or []:
                for row in table:
                    row_text = " | ".join(cell.strip() for cell in row if cell and cell.strip())
                    if row_text:
                        parts.append(row_text)
    return "\n\n".join(parts)


def _read_docx(path: str) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _split_into_sections(text: str) -> dict[str, str]:
    """Split full document text into named sections based on headings."""
    patterns = {
        "introduction": r"(?i)^\s*(\d+[\.\)]\s*)?introduction[:\.]?\s*$",
        "methods":      r"(?i)^\s*(\d+[\.\)]\s*)?methods?[:\.]?\s*$",
        "results":      r"(?i)^\s*(\d+[\.\)]\s*)?results?[:\.]?\s*$",
        "discussion":   r"(?i)^\s*(\d+[\.\)]\s*)?discussion[:\.]?\s*$",
    }
    sections: dict[str, str] = {}
    current, buf = None, []
    for line in text.split("\n"):
        stripped = line.strip()
        if len(stripped) <= 40:
            for name, pat in patterns.items():
                if re.match(pat, stripped):
                    if current:
                        sections[current] = "\n".join(buf).strip()
                    current, buf = name, []
                    break
            else:
                if current:
                    buf.append(line)
        elif current:
            buf.append(line)
    if current:
        sections[current] = "\n".join(buf).strip()
    return sections


def load_document(path: str) -> dict[str, str]:
    """Load a .pdf or .docx and return {section_name: section_text}."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _split_into_sections(_read_pdf(path))
    elif ext == ".docx":
        return _split_into_sections(_read_docx(path))
    raise ValueError(f"Unsupported format '{ext}'. Use .pdf or .docx.")


def build_full_paper_text(sections: dict[str, str]) -> str:
    """Combine detected sections into a single string for the LLM."""
    parts = [f"## {s.upper()}\n\n{sections[s]}" for s in SECTION_ORDER if sections.get(s)]
    return "\n\n---\n\n".join(parts)


# ── Token counting ────────────────────────────────────────────────────────────

_ENC        = tiktoken.get_encoding("o200k_base")
_MAX_TOKENS = 100_000


def check_token_count(text: str, label: str = "") -> int:
    n = len(_ENC.encode(text))
    tag = f"[{label}] " if label else ""
    if n > _MAX_TOKENS:
        raise ValueError(f"{tag}{n} tokens exceeds the {_MAX_TOKENS}-token limit.")
    if n > _MAX_TOKENS * 0.85:
        print(f"  ⚠  Warning: {tag}{n} tokens (>85% of limit).")
    return n


# ── Resource & rubric loading ─────────────────────────────────────────────────

def load_rubrics() -> dict[str, str]:
    """Load both rubric files into memory."""
    print("\nLoading rubrics...")
    rubrics = {}
    for name, path in [("original", RUBRIC_ORIGINAL_PATH), ("improved", RUBRIC_IMPROVED_PATH)]:
        if not os.path.exists(path):
            raise FileNotFoundError(f"Rubric not found: {path}")
        rubrics[name] = _read_docx(path)
        print(f"  ✓ '{name}' ({check_token_count(rubrics[name], label=name)} tokens)")
    return rubrics


def load_resources() -> dict[str, str]:
    """Load writing guide, sample results, and calibration summary into memory."""
    print("Loading resources...")
    resources = {}
    for label, path, reader in [
        ("writing_guide",       WRITING_GUIDE,       _read_pdf),
        ("sample_results",      SAMPLE_RESULTS,       _read_pdf),
        ("calibration_summary", CALIBRATION_SUMMARY, _read_docx),
    ]:
        if not os.path.exists(path):
            raise FileNotFoundError(f"Resource not found: {path}")
        resources[label] = reader(path)
        print(f"  ✓ '{label}' ({check_token_count(resources[label], label=label)} tokens)")
    return resources


# ── Pipeline & paper discovery ────────────────────────────────────────────────

def load_pipelines(path: str) -> list[dict]:
    """Load holistic pipelines from JSON."""
    with open(path, "r", encoding="utf-8") as f:
        all_pipelines = json.load(f)
    holistic = [p for p in all_pipelines if p.get("grading_mode") == "holistic"]
    print(f"Loaded {len(holistic)} holistic pipeline(s) from {len(all_pipelines)} total.")
    return holistic


def list_papers(folder: str) -> list[tuple[str, str]]:
    """Return sorted (student_id, path) pairs for all papers in folder."""
    paths = glob.glob(os.path.join(folder, "*.pdf")) + glob.glob(os.path.join(folder, "*.docx"))
    papers = [(os.path.splitext(os.path.basename(p))[0], p)
              for p in paths if not os.path.basename(p).startswith("~$")]
    return sorted(papers)


# ── Grade parsing & final grade computation ───────────────────────────────────

def _parse_grade(value) -> Optional[float]:
    """Accept any float in 1.0-10.0; snap to nearest 0.5. Returns None if invalid."""
    if value is None:
        return None
    try:
        g = float(value)
    except (TypeError, ValueError):
        return None
    if g < 1.0 or g > 10.0:
        return None
    return round(g * 2) / 2.0


def compute_final_grade(grades: dict[str, Optional[float]]) -> Optional[float]:
    """Weighted sum of subsection grades. Returns None if any grade is missing."""
    if any(grades.get(k) is None for k in WEIGHTS):
        return None
    return sum(grades[k] * w for k, w in WEIGHTS.items())


def parse_reply(raw: str) -> dict:
    """Parse the LLM JSON reply into a flat dict of grades, feedback, and rubric feedback."""
    empty = {k: None for k in [
        "intro_grade", "methods_grade", "results_grade", "discussion_grade", "lang_style_grade",
        "intro_feedback", "methods_feedback", "results_feedback",
        "discussion_feedback", "lang_style_feedback",
        "intro_rubric_feedback", "methods_rubric_feedback", "results_rubric_feedback",
        "discussion_rubric_feedback", "lang_style_rubric_feedback",
    ]}
    empty["raw_reply"] = raw
    try:
        data = json.loads(raw)
    except Exception:
        return empty
    return {
        "intro_grade":                _parse_grade(data.get("intro_grade")),
        "methods_grade":              _parse_grade(data.get("methods_grade")),
        "results_grade":              _parse_grade(data.get("results_grade")),
        "discussion_grade":           _parse_grade(data.get("discussion_grade")),
        "lang_style_grade":           _parse_grade(data.get("lang_style_grade")),
        "intro_feedback":             data.get("intro_feedback"),
        "methods_feedback":           data.get("methods_feedback"),
        "results_feedback":           data.get("results_feedback"),
        "discussion_feedback":        data.get("discussion_feedback"),
        "lang_style_feedback":        data.get("lang_style_feedback"),
        "intro_rubric_feedback":      data.get("intro_rubric_feedback"),
        "methods_rubric_feedback":    data.get("methods_rubric_feedback"),
        "results_rubric_feedback":    data.get("results_rubric_feedback"),
        "discussion_rubric_feedback": data.get("discussion_rubric_feedback"),
        "lang_style_rubric_feedback": data.get("lang_style_rubric_feedback"),
        "raw_reply":                  raw,
    }


# ── Prompt building ───────────────────────────────────────────────────────────

def build_system_prompt(rubric_text: str, resources: dict[str, str]) -> str:
    """Build the system prompt. Identical for both rubric versions."""
    return f"""{LLM_ROLE}

## Resources
The rubric is your primary grading tool — all grades must be grounded in its criteria. \
The other resources provide supporting context and do not override the rubric.

### Writing guide
{resources['writing_guide']}

### Sample Results section
{resources['sample_results']}

### Calibration summary
{resources['calibration_summary']}

## Grading rubric
{rubric_text}

## Rubric feedback
As you grade, provide feedback on each section of the rubric itself — not on the student's \
paper. Imagine you are developing and refining a grading rubric and are collecting notes to \
improve it. Do not describe or evaluate the student's work in these fields.

## Output format
Respond with a JSON object containing EXACTLY these keys:
{{
  "intro_grade": <number 1.0-10.0 in 0.5 steps>,
  "methods_grade": <number 1.0-10.0 in 0.5 steps>,
  "results_grade": <number 1.0-10.0 in 0.5 steps>,
  "discussion_grade": <number 1.0-10.0 in 0.5 steps>,
  "lang_style_grade": <number 1.0-10.0 in 0.5 steps>,
  "intro_feedback": "<reasoning behind the Introduction grade>",
  "methods_feedback": "<reasoning behind the Methods grade>",
  "results_feedback": "<reasoning behind the Results grade>",
  "discussion_feedback": "<reasoning behind the Discussion grade>",
  "lang_style_feedback": "<reasoning behind the Language & Style grade>",
  "intro_rubric_feedback": "<feedback on the Introduction rubric criteria, not the paper>",
  "methods_rubric_feedback": "<feedback on the Methods rubric criteria, not the paper>",
  "results_rubric_feedback": "<feedback on the Results rubric criteria, not the paper>",
  "discussion_rubric_feedback": "<feedback on the Discussion rubric criteria, not the paper>",
  "lang_style_rubric_feedback": "<feedback on the Language & Style rubric criteria, not the paper>"
}}""".strip()


# ── API call ──────────────────────────────────────────────────────────────────

def _call_openai(system_prompt: str, user_content: str, model: str, temperature: float) -> str:
    response = client.chat.completions.create(
        model=model,
        temperature=temperature,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": user_content},
        ],
        response_format={"type": "json_object"},
    )
    return response.choices[0].message.content


# ── CSV writing ───────────────────────────────────────────────────────────────

def append_result_row(csv_path: str, row: dict) -> None:
    """Append one row to the CSV, writing the header first if the file is new."""
    file_exists = os.path.exists(csv_path)
    os.makedirs(os.path.dirname(csv_path), exist_ok=True)
    with open(csv_path, "a", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=row.keys())
        if not file_exists:
            writer.writeheader()
        writer.writerow(row)


# ── Main grading function ─────────────────────────────────────────────────────

def grade_paper_holistic(
    student_id: str, sections: dict[str, str], rubric_text: str,
    rubric_version: str, model: str, temperature: float,
    run_id: str, resources: dict[str, str],
) -> dict:
    """Grade one paper in a single API call. Returns a CSV-ready row dict."""
    print(f"\n{'─' * 64}")
    print(f"  run_id : {run_id}")
    print(f"  model  : {model}  |  rubric: {rubric_version}  |  temp: {temperature}")
    print(f"{'─' * 64}")

    full_paper_text = build_full_paper_text(sections)
    check_token_count(full_paper_text, label=f"{student_id}.paper")

    system_prompt = build_system_prompt(rubric_text, resources)
    user_content  = "Please grade the following student paper.\n\n" + full_paper_text

    print(f"  System prompt : {len(system_prompt)} chars")
    print(f"  Paper preview : {full_paper_text[:120].replace(chr(10), ' ')} ...")

    raw_reply = _call_openai(system_prompt, user_content, model, temperature)
    parsed    = parse_reply(raw_reply)

    grades = {
        "introduction":   parsed["intro_grade"],
        "methods":        parsed["methods_grade"],
        "results":        parsed["results_grade"],
        "discussion":     parsed["discussion_grade"],
        "language_style": parsed["lang_style_grade"],
    }
    final_grade = compute_final_grade(grades)

    print(f"  intro={grades['introduction']}  methods={grades['methods']}  "
          f"results={grades['results']}  discussion={grades['discussion']}  "
          f"lang={grades['language_style']}  ->  final={final_grade}")

    if final_grade is None:
        print(f"  ⚠  Missing grade(s). Raw reply: {raw_reply[:200]}")

    return {
        "run_id":                     run_id,
        "student_id":                 student_id,
        "grading_mode":               "holistic",
        "rubric_version":             rubric_version,
        "model":                      model,
        "temperature":                temperature,
        "timestamp":                  datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "intro_grade":                parsed["intro_grade"],
        "intro_feedback":             parsed["intro_feedback"],
        "intro_rubric_feedback":      parsed["intro_rubric_feedback"],
        "methods_grade":              parsed["methods_grade"],
        "methods_feedback":           parsed["methods_feedback"],
        "methods_rubric_feedback":    parsed["methods_rubric_feedback"],
        "results_grade":              parsed["results_grade"],
        "results_feedback":           parsed["results_feedback"],
        "results_rubric_feedback":    parsed["results_rubric_feedback"],
        "discussion_grade":           parsed["discussion_grade"],
        "discussion_feedback":        parsed["discussion_feedback"],
        "discussion_rubric_feedback": parsed["discussion_rubric_feedback"],
        "lang_style_grade":           parsed["lang_style_grade"],
        "lang_style_feedback":        parsed["lang_style_feedback"],
        "lang_style_rubric_feedback": parsed["lang_style_rubric_feedback"],
        "final_grade":                final_grade,
    }


# ── Entry point ───────────────────────────────────────────────────────────────

if __name__ == "__main__":

    RUBRICS   = load_rubrics()
    RESOURCES = load_resources()
    PIPELINES = load_pipelines(PIPELINES_FILE)

    # Resume support — skip already completed run_ids
    completed: set[str] = set()
    if os.path.exists(RESULTS_CSV):
        with open(RESULTS_CSV, "r", encoding="utf-8") as f:
            completed = {row["run_id"] for row in csv.DictReader(f) if "run_id" in row}
    print(f"Skipping {len(completed)} already completed run(s).")

    papers     = list_papers(PAPERS_FOLDER)
    total_runs = len(papers) * len(PIPELINES) * N_RUNS
    done       = 0
    print(f"Found {len(papers)} paper(s): {[sid for sid, _ in papers]}\n")

    for student_id, doc_path in papers:
        try:
            sections = load_document(doc_path)
            print(f"\nPaper '{student_id}': sections -> {list(sections.keys())}")
            check_token_count(build_full_paper_text(sections), label=student_id)
        except Exception as e:
            print(f"  ✗ Cannot load '{student_id}': {e} — skipping.")
            continue

        for pipeline in PIPELINES:
            rubric_version = pipeline["rubric_version"]
            model          = pipeline["model"]
            temperature    = pipeline.get("temperature", 1.0)

            for run_idx in range(1, N_RUNS + 1):
                run_id = f"{student_id}.{pipeline['pipeline_id']}.run{run_idx}"

                if run_id in completed:
                    print(f"  -> Skipping {run_id} (already done).")
                    done += 1
                    continue

                try:
                    row = grade_paper_holistic(
                        student_id=student_id, sections=sections,
                        rubric_text=RUBRICS[rubric_version], rubric_version=rubric_version,
                        model=model, temperature=temperature,
                        run_id=run_id, resources=RESOURCES,
                    )
                    append_result_row(RESULTS_CSV, row)
                    done += 1
                    print(f"  ✓ Saved {run_id}  [{done}/{total_runs}]")
                except Exception as e:
                    print(f"  ✗ Error on {run_id}: {e}")

    print(f"\n{'=' * 64}")
    print(f"  Holistic grading complete. Results saved to: {RESULTS_CSV}")
    print(f"{'=' * 64}")