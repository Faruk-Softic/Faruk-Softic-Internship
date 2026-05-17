"""
Grades papers sequentially — one API call per section, plus one for Language & Style.
Each section receives relevant summaries from previously graded sections as context.
Final grade is computed in Python only — never by the LLM.
"""

import os
import csv
import json
import glob
import re
from datetime import datetime
from typing import Optional

import pdfplumber
import tiktoken
from docx import Document
from openai import OpenAI


# ── Paths & constants ─────────────────────────────────────────────────────────

API_KEY_FILE         = r"C:\Users\fsoft\Desktop\api_key.txt"
PAPERS_FOLDER        = r"C:\Users\fsoft\Desktop\Za Internship - code\Papers"
RUBRIC_ORIGINAL_PATH = r"C:\Users\fsoft\Desktop\Za Internship - code\Rubric_original.docx"
RUBRIC_IMPROVED_PATH = r"C:\Users\fsoft\Desktop\Za Internship - code\Rubric_improved.docx"
WRITING_GUIDE        = r"C:\Users\fsoft\Desktop\Za Internship - code\Writing_guide.pdf"
SAMPLE_RESULTS       = r"C:\Users\fsoft\Desktop\Za Internship - code\Sample_results.pdf"
CALIBRATION_SUMMARY  = r"C:\Users\fsoft\Desktop\Za Internship - code\Calibration_summary.docx"
PIPELINES_FILE       = r"C:\Users\fsoft\Desktop\Za Internship - code\pipelines.json"
RESULTS_CSV          = r"C:\Users\fsoft\Desktop\Za Internship - code\results.csv"

N_RUNS        = 5
SECTION_ORDER = ["introduction", "methods", "results", "discussion"]

WEIGHTS = {
    "introduction":   0.30,
    "methods":        0.15,
    "results":        0.15,
    "discussion":     0.30,
    "language_style": 0.10,
}

LLM_ROLE = (
    "You are a tutor grading first scientific papers written by second-year psychology "
    "students. The course focuses on scientific reasoning and argumentation. Students have "
    "only previously written a short literature review. Grade against the provided rubric, "
    "writing guide, and sample paper — not against publishable-paper standards. When "
    "deciding whether to fail a student, ask yourself whether they would be able to pass "
    "a bachelor's thesis the following year with the same skills."
)


# ── API client ────────────────────────────────────────────────────────────────

def _load_api_key() -> str:
    if not os.path.exists(API_KEY_FILE):
        raise FileNotFoundError(f"API key file not found: {API_KEY_FILE}")
    with open(API_KEY_FILE, "r", encoding="utf-8") as f:
        key = f.readline().strip()
    if not key:
        raise ValueError(f"API key file is empty: {API_KEY_FILE}")
    return key

client = OpenAI(api_key=_load_api_key())


# ── File readers ──────────────────────────────────────────────────────────────

def _read_pdf(path: str) -> str:
    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
            for table in page.extract_tables() or []:
                for row in table:
                    row_text = " | ".join(cell.strip() for cell in row if cell and cell.strip())
                    if row_text:
                        parts.append(row_text)
    return "\n\n".join(parts)


def _read_docx(path: str) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _split_into_sections(text: str) -> dict[str, str]:
    """Split full document text into named sections based on headings."""
    patterns = {
        "introduction": r"(?i)^\s*(\d+[\.\)]\s*)?introduction[:\.]?\s*$",
        "methods":      r"(?i)^\s*(\d+[\.\)]\s*)?methods?[:\.]?\s*$",
        "results":      r"(?i)^\s*(\d+[\.\)]\s*)?results?[:\.]?\s*$",
        "discussion":   r"(?i)^\s*(\d+[\.\)]\s*)?discussion[:\.]?\s*$",
    }
    sections: dict[str, str] = {}
    current, buf = None, []
    for line in text.split("\n"):
        stripped = line.strip()
        if len(stripped) <= 40:
            for name, pat in patterns.items():
                if re.match(pat, stripped):
                    if current:
                        sections[current] = "\n".join(buf).strip()
                    current, buf = name, []
                    break
            else:
                if current:
                    buf.append(line)
        elif current:
            buf.append(line)
    if current:
        sections[current] = "\n".join(buf).strip()
    return sections


def load_document(path: str) -> dict[str, str]:
    """Load a .pdf or .docx and return {section_name: section_text}."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _split_into_sections(_read_pdf(path))
    elif ext == ".docx":
        return _split_into_sections(_read_docx(path))
    raise ValueError(f"Unsupported format '{ext}'. Use .pdf or .docx.")


# ── Token counting ────────────────────────────────────────────────────────────

_ENC        = tiktoken.get_encoding("o200k_base")
_MAX_TOKENS = 100_000


def check_token_count(text: str, label: str = "") -> int:
    n = len(_ENC.encode(text))
    tag = f"[{label}] " if label else ""
    if n > _MAX_TOKENS:
        raise ValueError(f"{tag}{n} tokens exceeds the {_MAX_TOKENS}-token limit.")
    if n > _MAX_TOKENS * 0.85:
        print(f"  ⚠  Warning: {tag}{n} tokens (>85% of limit).")
    return n


# ── Resource & rubric loading ─────────────────────────────────────────────────

def load_rubrics() -> dict[str, str]:
    """Load both rubric files into memory."""
    print("\nLoading rubrics...")
    rubrics = {}
    for name, path in [("original", RUBRIC_ORIGINAL_PATH), ("improved", RUBRIC_IMPROVED_PATH)]:
        if not os.path.exists(path):
            raise FileNotFoundError(f"Rubric not found: {path}")
        rubrics[name] = _read_docx(path)
        print(f"  ✓ '{name}' ({check_token_count(rubrics[name], label=name)} tokens)")
    return rubrics


def load_resources() -> dict[str, str]:
    """Load writing guide, sample results, and calibration summary into memory."""
    print("Loading resources...")
    resources = {}
    for label, path, reader in [
        ("writing_guide",       WRITING_GUIDE,       _read_pdf),
        ("sample_results",      SAMPLE_RESULTS,       _read_pdf),
        ("calibration_summary", CALIBRATION_SUMMARY, _read_docx),
    ]:
        if not os.path.exists(path):
            raise FileNotFoundError(f"Resource not found: {path}")
        resources[label] = reader(path)
        print(f"  ✓ '{label}' ({check_token_count(resources[label], label=label)} tokens)")
    return resources


# ── Pipeline & paper discovery ────────────────────────────────────────────────

def load_pipelines(path: str) -> list[dict]:
    """Load sequential pipelines from JSON."""
    with open(path, "r", encoding="utf-8") as f:
        all_pipelines = json.load(f)
    sequential = [p for p in all_pipelines if p.get("grading_mode") == "sequential"]
    print(f"Loaded {len(sequential)} sequential pipeline(s) from {len(all_pipelines)} total.")
    return sequential


def list_papers(folder: str) -> list[tuple[str, str]]:
    """Return sorted (student_id, path) pairs for all papers in folder."""
    paths = glob.glob(os.path.join(folder, "*.pdf")) + glob.glob(os.path.join(folder, "*.docx"))
    papers = [(os.path.splitext(os.path.basename(p))[0], p)
              for p in paths if not os.path.basename(p).startswith("~$")]
    return sorted(papers)


# ── Grade parsing & final grade computation ───────────────────────────────────

def _parse_grade(value) -> Optional[float]:
    """Accept any float in 1.0-10.0; snap to nearest 0.5. Returns None if invalid."""
    if value is None:
        return None
    try:
        g = float(value)
    except (TypeError, ValueError):
        return None
    if g < 1.0 or g > 10.0:
        return None
    return round(g * 2) / 2.0


def compute_final_grade(grades: dict[str, Optional[float]]) -> Optional[float]:
    """Weighted sum of subsection grades. Returns None if any grade is missing."""
    if any(grades.get(k) is None for k in WEIGHTS):
        return None
    return sum(grades[k] * w for k, w in WEIGHTS.items())


def _parse_section_reply(raw: str) -> dict:
    """Parse a section grading reply into grades, feedback, and summaries."""
    empty = {k: None for k in [
        "grade", "feedback", "content_summary", "style_notes", "rubric_feedback"
    ]}
    empty["raw_reply"] = raw
    try:
        data = json.loads(raw)
    except Exception:
        return empty
    return {
        "grade":            _parse_grade(data.get("grade")),
        "feedback":         data.get("feedback"),
        "content_summary":  data.get("content_summary"),
        "style_notes":      data.get("style_notes"),
        "rubric_feedback":  data.get("rubric_feedback"),
        "raw_reply":        raw,
    }


def _parse_lang_style_reply(raw: str) -> dict:
    """Parse the Language & Style grading reply."""
    empty = {"grade": None, "feedback": None, "rubric_feedback": None, "raw_reply": raw}
    try:
        data = json.loads(raw)
    except Exception:
        return empty
    return {
        "grade":           _parse_grade(data.get("grade")),
        "feedback":        data.get("feedback"),
        "rubric_feedback": data.get("rubric_feedback"),
        "raw_reply":       raw,
    }


# ── Prompt building ───────────────────────────────────────────────────────────

def _build_resource_block(resources: dict[str, str]) -> str:
    return (
        "### Writing guide\n"       f"{resources['writing_guide']}\n\n"
        "### Sample Results section\n" f"{resources['sample_results']}\n\n"
        "### Calibration summary\n" f"{resources['calibration_summary']}\n"
    )


# Context passed forward between sections:
#   introduction → (none)
#   methods      → introduction content_summary
#   results      → methods content_summary
#   discussion   → introduction content_summary + style_notes, results content_summary
_CONTEXT_RULES = {
    "introduction": [],
    "methods":      [("introduction", ["content_summary"])],
    "results":      [("methods",      ["content_summary"])],
    "discussion":   [("introduction", ["content_summary", "style_notes"]),
                     ("results",      ["content_summary"])],
}


def build_section_prompt(
    section_name: str, rubric_text: str,
    previous_summaries: dict[str, dict], resources: dict[str, str],
) -> str:
    """Build the system prompt for grading one section."""
    # Context block from previous sections
    context_parts = []
    for src_sec, fields in _CONTEXT_RULES[section_name]:
        if src_sec in previous_summaries:
            s = previous_summaries[src_sec]
            for field in fields:
                label = field.replace("_", " ").capitalize()
                context_parts.append(f"**{src_sec.capitalize()} — {label}:**\n{s.get(field, '')}")
    context_block = (
        "## Context from previously graded sections\n\n" + "\n\n".join(context_parts)
        if context_parts else ""
    )

    return f"""{LLM_ROLE}

## Resources
The rubric is your primary grading tool. The resources below provide supporting context \
and do not override the rubric.

{_build_resource_block(resources)}

{context_block}

## Grading rubric
{rubric_text}

## Rubric feedback
As you grade, provide feedback on this section's rubric criteria themselves — not on the \
student's paper. Imagine you are developing and refining a grading rubric and are collecting \
notes to improve it. Do not describe or evaluate the student's work in this field.

## Style notes
Write brief style notes covering: clarity and precision of language, grammar and spelling, \
APA formatting, paragraph structure, and any redundancy or inconsistency. These will be \
used later to assign an overall Language & Style grade.

## Output format
Respond with a JSON object containing EXACTLY these keys:
{{
  "grade": <number 1.0-10.0 in 0.5 steps>,
  "feedback": "<reasoning behind the grade for this section>",
  "content_summary": "<3-5 sentence summary of the scientific content and quality of this section>",
  "style_notes": "<2-4 sentence summary of language and style observations>",
  "rubric_feedback": "<feedback on this section's rubric criteria, not the paper>"
}}""".strip()


def build_lang_style_prompt(
    rubric_text: str, all_style_notes: dict[str, str], resources: dict[str, str],
) -> str:
    """Build the system prompt for the Language & Style grading step."""
    notes_block = "\n\n".join(
        f"**{sec.capitalize()} — style notes:**\n{note}"
        for sec, note in all_style_notes.items() if note
    )
    return f"""{LLM_ROLE}

## Resources
The rubric is your primary grading tool. The resources below provide supporting context \
and do not override the rubric.

{_build_resource_block(resources)}

## Style notes from all sections
{notes_block}

## Grading rubric
{rubric_text}

## Rubric feedback
As you grade, provide feedback on the Language & Style rubric criteria themselves — not on \
the student's paper. Imagine you are developing and refining a grading rubric and are \
collecting notes to improve it. Do not describe or evaluate the student's work in this field.

## Output format
Respond with a JSON object containing EXACTLY these keys:
{{
  "grade": <number 1.0-10.0 in 0.5 steps>,
  "feedback": "<reasoning behind the Language & Style grade>",
  "rubric_feedback": "<feedback on the Language & Style rubric criteria, not the paper>"
}}""".strip()


# ── API call ──────────────────────────────────────────────────────────────────

def _call_openai(system_prompt: str, user_content: str, model: str, temperature: float) -> str:
    response = client.chat.completions.create(
        model=model,
        temperature=temperature,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": user_content},
        ],
        response_format={"type": "json_object"},
    )
    return response.choices[0].message.content


# ── CSV writing ───────────────────────────────────────────────────────────────

def append_result_row(csv_path: str, row: dict) -> None:
    """Append one row to the CSV, writing the header first if the file is new."""
    file_exists = os.path.exists(csv_path)
    os.makedirs(os.path.dirname(csv_path), exist_ok=True)
    with open(csv_path, "a", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=row.keys())
        if not file_exists:
            writer.writeheader()
        writer.writerow(row)


# ── Main grading function ─────────────────────────────────────────────────────

def grade_paper_sequential(
    student_id: str, sections: dict[str, str], rubric_text: str,
    rubric_version: str, model: str, temperature: float,
    run_id: str, resources: dict[str, str],
) -> dict:
    """Grade one paper sequentially. Returns a CSV-ready row dict."""
    print(f"\n{'─' * 64}")
    print(f"  run_id : {run_id}")
    print(f"  model  : {model}  |  rubric: {rubric_version}  |  temp: {temperature}")
    print(f"{'─' * 64}")

    previous_summaries: dict[str, dict] = {}
    section_results:    dict[str, dict] = {}

    # Grade each section sequentially
    for sec in SECTION_ORDER:
        if sec not in sections:
            print(f"  ⚠  '{sec}' not found — skipping.")
            section_results[sec] = {k: None for k in
                ["grade", "feedback", "content_summary", "style_notes", "rubric_feedback", "raw_reply"]}
            continue

        system_prompt = build_section_prompt(sec, rubric_text, previous_summaries, resources)
        user_content  = f"Please grade the following {sec.upper()} section:\n\n{sections[sec]}"
        print(f"  System prompt [{sec}] : {len(system_prompt)} chars")

        raw    = _call_openai(system_prompt, user_content, model, temperature)
        result = _parse_section_reply(raw)
        section_results[sec] = result

        print(f"  ✓ {sec.upper()}: grade={result['grade']}  "
              f"feedback={str(result['feedback'] or '')[:80]} ...")

        if result["grade"] is None:
            print(f"  ⚠  Could not parse grade for {sec}. Raw: {raw[:200]}")

        previous_summaries[sec] = {
            "content_summary": result["content_summary"],
            "style_notes":     result["style_notes"],
        }

    # Language & Style
    all_style_notes = {s: section_results[s]["style_notes"]
                       for s in SECTION_ORDER if section_results[s]["style_notes"]}
    ls_prompt = build_lang_style_prompt(rubric_text, all_style_notes, resources)
    ls_user   = "Using the style notes above and the rubric, assign an overall Language & Style grade."
    print(f"  System prompt [lang_style] : {len(ls_prompt)} chars")

    ls_raw    = _call_openai(ls_prompt, ls_user, model, temperature)
    ls_result = _parse_lang_style_reply(ls_raw)

    print(f"  ✓ LANG & STYLE: grade={ls_result['grade']}  "
          f"feedback={str(ls_result['feedback'] or '')[:80]} ...")

    if ls_result["grade"] is None:
        print(f"  ⚠  Could not parse Language & Style grade. Raw: {ls_raw[:200]}")

    # Final grade (Python only)
    grades = {
        "introduction":   section_results["introduction"]["grade"],
        "methods":        section_results["methods"]["grade"],
        "results":        section_results["results"]["grade"],
        "discussion":     section_results["discussion"]["grade"],
        "language_style": ls_result["grade"],
    }
    final_grade = compute_final_grade(grades)
    print(f"  -> final={final_grade}  (weighted, Python)")

    if final_grade is None:
        print("  ⚠  Could not compute final grade — one or more component grades are missing.")

    # Build CSV row
    r = section_results
    return {
        "run_id":                     run_id,
        "student_id":                 student_id,
        "grading_mode":               "sequential",
        "rubric_version":             rubric_version,
        "model":                      model,
        "temperature":                temperature,
        "timestamp":                  datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "intro_grade":                r["introduction"]["grade"],
        "intro_feedback":             r["introduction"]["feedback"],
        "intro_content_summary":      r["introduction"]["content_summary"],
        "intro_style_notes":          r["introduction"]["style_notes"],
        "intro_rubric_feedback":      r["introduction"]["rubric_feedback"],
        "methods_grade":              r["methods"]["grade"],
        "methods_feedback":           r["methods"]["feedback"],
        "methods_content_summary":    r["methods"]["content_summary"],
        "methods_style_notes":        r["methods"]["style_notes"],
        "methods_rubric_feedback":    r["methods"]["rubric_feedback"],
        "results_grade":              r["results"]["grade"],
        "results_feedback":           r["results"]["feedback"],
        "results_content_summary":    r["results"]["content_summary"],
        "results_style_notes":        r["results"]["style_notes"],
        "results_rubric_feedback":    r["results"]["rubric_feedback"],
        "discussion_grade":           r["discussion"]["grade"],
        "discussion_feedback":        r["discussion"]["feedback"],
        "discussion_content_summary": r["discussion"]["content_summary"],
        "discussion_style_notes":     r["discussion"]["style_notes"],
        "discussion_rubric_feedback": r["discussion"]["rubric_feedback"],
        "lang_style_grade":           ls_result["grade"],
        "lang_style_feedback":        ls_result["feedback"],
        "lang_style_rubric_feedback": ls_result["rubric_feedback"],
        "final_grade":                final_grade,
    }


# ── Entry point ───────────────────────────────────────────────────────────────

if __name__ == "__main__":

    RUBRICS   = load_rubrics()
    RESOURCES = load_resources()
    PIPELINES = load_pipelines(PIPELINES_FILE)

    # Resume support — skip already completed run_ids
    completed: set[str] = set()
    if os.path.exists(RESULTS_CSV):
        with open(RESULTS_CSV, "r", encoding="utf-8") as f:
            completed = {row["run_id"] for row in csv.DictReader(f) if "run_id" in row}
    print(f"Skipping {len(completed)} already completed run(s).")

    papers     = list_papers(PAPERS_FOLDER)
    total_runs = len(papers) * len(PIPELINES) * N_RUNS
    done       = 0
    print(f"Found {len(papers)} paper(s): {[sid for sid, _ in papers]}\n")

    for student_id, doc_path in papers:
        try:
            sections = load_document(doc_path)
            print(f"\nPaper '{student_id}': sections -> {list(sections.keys())}")
            for sec, text in sections.items():
                check_token_count(text, label=f"{student_id}.{sec}")
        except Exception as e:
            print(f"  ✗ Cannot load '{student_id}': {e} — skipping.")
            continue

        for pipeline in PIPELINES:
            rubric_version = pipeline["rubric_version"]
            model          = pipeline["model"]
            temperature    = pipeline.get("temperature", 1.0)

            for run_idx in range(1, N_RUNS + 1):
                run_id = f"{student_id}.{pipeline['pipeline_id']}.run{run_idx}"

                if run_id in completed:
                    print(f"  -> Skipping {run_id} (already done).")
                    done += 1
                    continue

                try:
                    row = grade_paper_sequential(
                        student_id=student_id, sections=sections,
                        rubric_text=RUBRICS[rubric_version], rubric_version=rubric_version,
                        model=model, temperature=temperature,
                        run_id=run_id, resources=RESOURCES,
                    )
                    append_result_row(RESULTS_CSV, row)
                    done += 1
                    print(f"  ✓ Saved {run_id}  [{done}/{total_runs}]")
                except Exception as e:
                    print(f"  ✗ Error on {run_id}: {e}")

    print(f"\n{'=' * 64}")
    print(f"  Sequential grading complete. Results saved to: {RESULTS_CSV}")
    print(f"{'=' * 64}")