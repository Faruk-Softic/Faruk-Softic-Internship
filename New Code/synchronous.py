"""
Grades papers holistically in a single API call per run.
Checklist items (0/1), section grades, and feedback are all written to CSV.
Final grade is computed in Python only — never by the LLM.
"""

import os, csv, json, glob, re
from datetime import datetime
from typing import Optional

import pdfplumber, tiktoken
from docx import Document
from openai import OpenAI
from dotenv import load_dotenv

load_dotenv()

# ── Paths & constants ─────────────────────────────────────────────────────────

PAPERS_FOLDER        = os.environ.get("PAPERS_FOLDER")
RUBRIC_ORIGINAL_PATH = os.environ.get("RUBRIC_ORIGINAL_PATH")
RUBRIC_IMPROVED_PATH = os.environ.get("RUBRIC_IMPROVED_PATH")
WRITING_GUIDE        = os.environ.get("WRITING_GUIDE")
SAMPLE_RESULTS       = os.environ.get("SAMPLE_RESULTS")
CALIBRATION_SUMMARY  = os.environ.get("CALIBRATION_SUMMARY")
PIPELINES_FILE       = os.environ.get("PIPELINES_FILE")
RESULTS_CSV          = os.environ.get("RESULTS_CSV")
PROMPTS_FILE         = os.environ.get("PROMPTS_FILE")

N_RUNS = 1

WEIGHTS = {
    "introduction":   0.30,
    "methods":        0.15,
    "results":        0.15,
    "discussion":     0.30,
    "language_style": 0.10,
}

TIER_ORDER  = ["tier0", "tier1", "tier2", "tier3", "tier4"]
ITEM_TYPES  = {
    "tier0": ["ko"],
    "tier1": ["core_criteria", "additional_criteria"],
    "tier2": ["core_criteria", "additional_criteria"],
    "tier3": ["core_criteria", "additional_criteria"],
    "tier4": ["core_criteria", "additional_criteria"],
}
TYPE_ABBREV = {"ko": "ko", "core_criteria": "cc", "additional_criteria": "ac"}
TYPE_LABELS = {
    "ko":                  "KNOCK-OUT — paper fails to reach Tier 1 if ANY apply:",
    "core_criteria":       "CORE CRITERIA — all must be met to reach this tier:",
    "additional_criteria": "ADDITIONAL CRITERIA — push grade toward top of tier:",
}
TIER_LABELS = {
    "tier0": "Tier 0 — Insufficient (1.0–5.0)",
    "tier1": "Tier 1 — Passing (5.5–6.4)",
    "tier2": "Tier 2 — Adequate (6.5–7.4)",
    "tier3": "Tier 3 — Good (7.5–8.4)",
    "tier4": "Tier 4 — Excellent (8.5–10.0)",
}
SECTION_LABELS = {
    "introduction":   "Introduction",
    "methods":        "Methods",
    "results":        "Results",
    "discussion":     "Discussion",
    "language_style": "Language & Style",
}

client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))


# ── Prompt loading & building ─────────────────────────────────────────────────

def load_prompts(path: str) -> dict:
    """Load all prompt strings and checklist data from prompts.json."""
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def get_checklist_keys(checklist: dict) -> list[str]:
    """
    Return an ordered list of all checklist column names, e.g.:
      introduction_t0_ko_1, introduction_t1_cc_1, introduction_t1_ac_1, ...
      language_style_t4_ac_2
    """
    keys = []
    for section, tiers in checklist.items():
        for tier in TIER_ORDER:
            if tier not in tiers:
                continue
            t = tier.replace("tier", "t")
            for item_type in ITEM_TYPES[tier]:
                if item_type not in tiers[tier]:
                    continue
                abbrev = TYPE_ABBREV[item_type]
                for idx in range(1, len(tiers[tier][item_type]) + 1):
                    keys.append(f"{section}_{t}_{abbrev}_{idx}")
    return keys


def build_rubric_text(checklist: dict) -> str:
    """Render the structured checklist into the plain-text rubric the LLM sees."""
    lines = [
        "Tiers are cumulative. A paper must satisfy all CORE CRITERIA items in a tier "
        "(and all tiers below it) to be graded within that tier's range. "
        "ADDITIONAL CRITERIA items push the grade toward the top of the range but are not required.\n"
        "Tier 0 items are knock-out criteria: if ANY apply, the paper cannot reach Tier 1.\n\n"
        "Tier ranges:\n"
        "  Tier 0 — Insufficient  : 1.0 – 5.0\n"
        "  Tier 1 — Passing       : 5.5 – 6.4\n"
        "  Tier 2 — Adequate      : 6.5 – 7.4\n"
        "  Tier 3 — Good          : 7.5 – 8.4\n"
        "  Tier 4 — Excellent     : 8.5 – 10.0"
    ]
    for section, tiers in checklist.items():
        sec_label = SECTION_LABELS.get(section, section)
        lines += [f"\n{'=' * 60}", f"SECTION: {sec_label.upper()}", "=" * 60]
        for tier in TIER_ORDER:
            if tier not in tiers:
                continue
            lines += [f"\n{TIER_LABELS[tier]}", "-" * 40]
            for item_type in ITEM_TYPES[tier]:
                if item_type not in tiers[tier]:
                    continue
                lines.append(TYPE_LABELS[item_type])
                t      = tier.replace("tier", "t")
                abbrev = TYPE_ABBREV[item_type]
                for idx, text in enumerate(tiers[tier][item_type], start=1):
                    key = f"{section}_{t}_{abbrev}_{idx}"
                    lines.append(f"  [{key}] {text}")
    return "\n".join(lines)


def build_output_format(checklist: dict, prompts: dict) -> str:
    """
    Generate the full output-format specification shown to the LLM,
    with every checklist key and all grade/feedback keys listed explicitly.
    """
    lines = [prompts["output_format_preamble"], "", "{"]
    for section, tiers in checklist.items():
        lines.append(f'  // --- {SECTION_LABELS.get(section, section)} checklist ---')
        for tier in TIER_ORDER:
            if tier not in tiers:
                continue
            t = tier.replace("tier", "t")
            for item_type in ITEM_TYPES[tier]:
                if item_type not in tiers[tier]:
                    continue
                abbrev = TYPE_ABBREV[item_type]
                for idx in range(1, len(tiers[tier][item_type]) + 1):
                    key = f"{section}_{t}_{abbrev}_{idx}"
                    lines.append(f'  "{key}": <0 or 1>,')
        lines.append("")

    grade_feedback_keys = [
        ("introduction",   "intro_grade",      "intro_feedback"),
        ("methods",        "methods_grade",     "methods_feedback"),
        ("results",        "results_grade",     "results_feedback"),
        ("discussion",     "discussion_grade",  "discussion_feedback"),
        ("language_style", "lang_style_grade",  "lang_style_feedback"),
    ]
    lines.append("  // --- Section grades and feedback ---")
    for section, grade_key, feedback_key in grade_feedback_keys:
        label = SECTION_LABELS.get(section, section)
        lines.append(f'  "{grade_key}": <number 1.0-10.0 in 0.5 steps, or null if section absent>,')
        lines.append(f'  "{feedback_key}": "<reasoning behind the {label} grade>",')

    lines[-1] = lines[-1].rstrip(",")
    lines.append("}")
    return "\n".join(lines)


def build_system_prompt(rubric_text: str, resources: dict, prompts: dict,
                        output_format: str) -> str:
    return prompts["system_prompt_template"].format(
        llm_role            = prompts["llm_role"],
        resources_preamble  = prompts["resources_preamble"],
        writing_guide       = resources["writing_guide"],
        sample_results      = resources["sample_results"],
        calibration_summary = resources["calibration_summary"],
        grading_notes       = prompts["grading_notes"],
        rubric              = rubric_text,
        output_format       = output_format,
    ).strip()


def build_user_message(paper_text: str, prompts: dict) -> str:
    return prompts["user_message_prefix"] + paper_text + prompts["user_message_suffix"]


# ── File readers ──────────────────────────────────────────────────────────────

def _read_pdf(path: str) -> str:
    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
            for table in page.extract_tables() or []:
                for row in table:
                    row_text = " | ".join(cell.strip() for cell in row if cell and cell.strip())
                    if row_text:
                        parts.append(row_text)
    return "\n\n".join(parts)


def _read_docx(path: str) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _read_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def _split_into_sections(text: str) -> dict[str, str]:
    patterns = {
        "introduction": r"(?i)^\s*(\d+[\.\)]\s*)?(introduction|introductie)[:\.]?\s*$",
        "methods":      r"(?i)^\s*(\d+[\.\)]\s*)?(methods?|methoden?|methode?|methodologie)[:\.]?\s*$",
        "results":      r"(?i)^\s*(\d+[\.\)]\s*)?(results?|resultaten)[:\.]?\s*$",
        "discussion":   r"(?i)^\s*(\d+[\.\)]\s*)?(discussion|discussie)[:\.]?\s*$",
    }
    sections: dict[str, str] = {}
    current, buf = None, []
    for line in text.split("\n"):
        stripped = line.strip()
        if len(stripped) <= 40:
            for name, pat in patterns.items():
                if re.match(pat, stripped):
                    if current:
                        sections[current] = "\n".join(buf).strip()
                    current, buf = name, []
                    break
            else:
                if current:
                    buf.append(line)
        elif current:
            buf.append(line)
    if current:
        sections[current] = "\n".join(buf).strip()
    return sections


def load_document(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":  return _read_pdf(path)
    if ext == ".docx": return _read_docx(path)
    if ext == ".txt":  return _read_txt(path)
    raise ValueError(f"Unsupported format '{ext}'. Use .pdf, .docx, or .txt.")


# ── Token counting ────────────────────────────────────────────────────────────

_ENC        = tiktoken.get_encoding("o200k_base")
_MAX_TOKENS = 100_000


def check_token_count(text: str, label: str = "") -> int:
    n   = len(_ENC.encode(text))
    tag = f"[{label}] " if label else ""
    if n > _MAX_TOKENS:
        raise ValueError(f"{tag}{n} tokens exceeds the {_MAX_TOKENS}-token limit.")
    if n > _MAX_TOKENS * 0.85:
        print(f"  ⚠  Warning: {tag}{n} tokens (>85% of limit).")
    return n


# ── Resource & rubric loading ─────────────────────────────────────────────────

def load_rubrics() -> dict[str, str]:
    print("\nLoading rubrics...")
    rubrics = {}
    for name, path in [("original", RUBRIC_ORIGINAL_PATH), ("improved", RUBRIC_IMPROVED_PATH)]:
        if not os.path.exists(path):
            raise FileNotFoundError(f"Rubric not found: {path}")
        rubrics[name] = _read_docx(path)
        print(f"  ✓ '{name}' ({check_token_count(rubrics[name], label=name)} tokens)")
    return rubrics


def load_resources() -> dict[str, str]:
    print("Loading resources...")
    resources = {}
    for label, path, reader in [
        ("writing_guide",       WRITING_GUIDE,       _read_pdf),
        ("sample_results",      SAMPLE_RESULTS,      _read_pdf),
        ("calibration_summary", CALIBRATION_SUMMARY, _read_docx),
    ]:
        if not os.path.exists(path):
            raise FileNotFoundError(f"Resource not found: {path}")
        resources[label] = reader(path)
        print(f"  ✓ '{label}' ({check_token_count(resources[label], label=label)} tokens)")
    return resources


# ── Pipeline & paper discovery ────────────────────────────────────────────────

def load_pipelines(path: str) -> list[dict]:
    with open(path, "r", encoding="utf-8") as f:
        all_pipelines = json.load(f)
    holistic = [p for p in all_pipelines if p.get("grading_mode") == "holistic"]
    print(f"Loaded {len(holistic)} holistic pipeline(s) from {len(all_pipelines)} total.")
    return holistic


def list_papers(folder: str) -> list[tuple[str, str]]:
    paths = (glob.glob(os.path.join(folder, "*.pdf")) +
             glob.glob(os.path.join(folder, "*.docx")) +
             glob.glob(os.path.join(folder, "*.txt")))
    papers = [(os.path.splitext(os.path.basename(p))[0], p)
              for p in paths if not os.path.basename(p).startswith("~$")]
    return sorted(papers)


# ── Grade parsing & final grade computation ───────────────────────────────────

def _parse_grade(value) -> Optional[float]:
    if value is None:
        return None
    try:
        g = float(value)
    except (TypeError, ValueError):
        return None
    if g < 1.0 or g > 10.0:
        return None
    return round(g * 2) / 2.0


def _parse_binary(value) -> Optional[int]:
    """Accept 0 or 1 only. Returns None if value is missing or invalid."""
    if value is None:
        return None
    try:
        v = int(value)
    except (TypeError, ValueError):
        return None
    return v if v in (0, 1) else None


def compute_final_grade(grades: dict[str, Optional[float]]) -> Optional[float]:
    if any(grades.get(k) is None for k in WEIGHTS):
        return None
    return sum(grades[k] * w for k, w in WEIGHTS.items())


def parse_reply(raw: str, checklist_keys: list[str]) -> dict:
    """Parse the LLM JSON reply into checklist 0/1 values, grades, and feedback."""
    result = {k: None for k in checklist_keys}
    for k in ["intro_grade", "methods_grade", "results_grade", "discussion_grade",
              "lang_style_grade", "intro_feedback", "methods_feedback",
              "results_feedback", "discussion_feedback", "lang_style_feedback"]:
        result[k] = None
    result["raw_reply"]        = raw
    result["missing_sections"] = ""

    try:
        data = json.loads(raw)
    except Exception:
        return result

    for key in checklist_keys:
        result[key] = _parse_binary(data.get(key))

    for k in ["intro_grade", "methods_grade", "results_grade",
              "discussion_grade", "lang_style_grade"]:
        result[k] = _parse_grade(data.get(k))

    for k in ["intro_feedback", "methods_feedback", "results_feedback",
              "discussion_feedback", "lang_style_feedback"]:
        result[k] = data.get(k)

    missing = [
        sec for sec, gk in {
            "introduction":   "intro_grade",
            "methods":        "methods_grade",
            "results":        "results_grade",
            "discussion":     "discussion_grade",
            "language_style": "lang_style_grade",
        }.items() if result[gk] is None
    ]
    result["missing_sections"] = ", ".join(missing) if missing else ""
    return result


# ── API call ──────────────────────────────────────────────────────────────────

def _call_openai(system_prompt: str, user_content: str,
                 model: str, temperature: float) -> str:
    response = client.responses.create(
        model=model,
        temperature=temperature,
        instructions=system_prompt,
        input=user_content,
        text={"format": {"type": "json_object"}},
    )
    return response.output_text


# ── CSV writing ───────────────────────────────────────────────────────────────

def append_result_row(csv_path: str, row: dict) -> None:
    file_exists = os.path.exists(csv_path)
    os.makedirs(os.path.dirname(csv_path), exist_ok=True)
    with open(csv_path, "a", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=row.keys())
        if not file_exists:
            writer.writeheader()
        writer.writerow(row)


# ── Main grading function ─────────────────────────────────────────────────────

def grade_paper_holistic(
    student_id: str, paper_text: str, rubric_text: str,
    rubric_version: str, model: str, temperature: float,
    run_id: str, resources: dict, prompts: dict,
    checklist_keys: list[str], output_format: str,
) -> dict:
    print(f"\n{'─' * 64}")
    print(f"  run_id : {run_id}")
    print(f"  model  : {model}  |  rubric: {rubric_version}  |  temp: {temperature}")
    print(f"{'─' * 64}")

    system_prompt = build_system_prompt(rubric_text, resources, prompts, output_format)
    user_content  = build_user_message(paper_text, prompts)

    print(f"  System prompt : {len(system_prompt)} chars")
    print(f"  Paper preview : {paper_text[:120].replace(chr(10), ' ')} ...")

    raw_reply = _call_openai(system_prompt, user_content, model, temperature)
    parsed    = parse_reply(raw_reply, checklist_keys)

    grades = {
        "introduction":   parsed["intro_grade"],
        "methods":        parsed["methods_grade"],
        "results":        parsed["results_grade"],
        "discussion":     parsed["discussion_grade"],
        "language_style": parsed["lang_style_grade"],
    }
    final_grade = compute_final_grade(grades)

    print(f"  intro={grades['introduction']}  methods={grades['methods']}  "
          f"results={grades['results']}  discussion={grades['discussion']}  "
          f"lang={grades['language_style']}  ->  final={final_grade}")

    if parsed["missing_sections"]:
        print(f"  ⚠  Missing section(s): {parsed['missing_sections']}")
    if final_grade is None:
        print(f"  ⚠  Missing grade(s). Raw reply: {raw_reply[:200]}")

    # ── CSV row: metadata | checklist 0/1 | grades + feedback | final ─────────
    row = {
        "run_id":         run_id,
        "student_id":     student_id,
        "grading_mode":   "holistic",
        "rubric_version": rubric_version,
        "model":          model,
        "temperature":    temperature,
        "timestamp":      datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }
    for key in checklist_keys:
        row[key] = parsed[key]
    row.update({
        "intro_grade":         parsed["intro_grade"],
        "intro_feedback":      parsed["intro_feedback"],
        "methods_grade":       parsed["methods_grade"],
        "methods_feedback":    parsed["methods_feedback"],
        "results_grade":       parsed["results_grade"],
        "results_feedback":    parsed["results_feedback"],
        "discussion_grade":    parsed["discussion_grade"],
        "discussion_feedback": parsed["discussion_feedback"],
        "lang_style_grade":    parsed["lang_style_grade"],
        "lang_style_feedback": parsed["lang_style_feedback"],
        "final_grade":         final_grade,
        "missing_sections":    parsed["missing_sections"],
    })
    return row


# ── Entry point ───────────────────────────────────────────────────────────────

if __name__ == "__main__":

    PROMPTS   = load_prompts(PROMPTS_FILE)
    RUBRICS   = load_rubrics()
    RESOURCES = load_resources()
    PIPELINES = load_pipelines(PIPELINES_FILE)

    checklist      = PROMPTS["checklist"]
    CHECKLIST_KEYS = get_checklist_keys(checklist)
    RUBRIC_TEXT    = build_rubric_text(checklist)
    OUTPUT_FORMAT  = build_output_format(checklist, PROMPTS)

    print(f"Checklist columns : {len(CHECKLIST_KEYS)}")

    completed: set[str] = set()
    if os.path.exists(RESULTS_CSV):
        with open(RESULTS_CSV, "r", encoding="utf-8") as f:
            completed = {row["run_id"] for row in csv.DictReader(f) if "run_id" in row}
    print(f"Skipping {len(completed)} already completed run(s).")

    papers     = list_papers(PAPERS_FOLDER)
    total_runs = len(papers) * len(PIPELINES) * N_RUNS
    done       = 0
    print(f"Found {len(papers)} paper(s): {[sid for sid, _ in papers]}\n")

    for student_id, doc_path in papers:
        try:
            paper_text = load_document(doc_path)
            print(f"\nPaper '{student_id}': {len(paper_text)} characters loaded.")
            check_token_count(paper_text, label=student_id)
        except Exception as e:
            print(f"  ✗ Cannot load '{student_id}': {e} — skipping.")
            continue

        for pipeline in PIPELINES:
            rubric_version = pipeline["rubric_version"]
            model          = pipeline["model"]
            temperature    = pipeline.get("temperature", 1.0)

            # Improved rubric: auto-generated from structured checklist.
            # Original rubric: loaded from file as before.
            rubric_text = (RUBRIC_TEXT if rubric_version == "improved"
                           else RUBRICS[rubric_version])

            for run_idx in range(1, N_RUNS + 1):
                run_id = f"{student_id}.{pipeline['pipeline_id']}.run{run_idx}"

                if run_id in completed:
                    print(f"  -> Skipping {run_id} (already done).")
                    done += 1
                    continue

                try:
                    row = grade_paper_holistic(
                        student_id=student_id, paper_text=paper_text,
                        rubric_text=rubric_text, rubric_version=rubric_version,
                        model=model, temperature=temperature,
                        run_id=run_id, resources=RESOURCES, prompts=PROMPTS,
                        checklist_keys=CHECKLIST_KEYS, output_format=OUTPUT_FORMAT,
                    )
                    append_result_row(RESULTS_CSV, row)
                    done += 1
                    print(f"  ✓ Saved {run_id}  [{done}/{total_runs}]")
                except Exception as e:
                    print(f"  ✗ Error on {run_id}: {e}")

    print(f"\n{'=' * 64}")
    print(f"  Holistic grading complete. Results saved to: {RESULTS_CSV}")
    print(f"{'=' * 64}")
